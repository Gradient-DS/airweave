"""Simple DOCX text extraction using python-docx."""

import asyncio
from pathlib import Path
from typing import Dict, List

from docx import Document

from airweave.core.async_helpers import run_in_thread_pool
from airweave.core.exceptions import EntityProcessingError
from airweave.platform.converters._base import BaseTextConverter


class SimpleDocxConverter(BaseTextConverter):
    """Extract text from DOCX files using python-docx (local, no API)."""

    BATCH_SIZE = 10
    MAX_CONCURRENT = 20

    def __init__(self):
        """Initialize SimpleDocxConverter."""
        super().__init__()
        self.logger.info("Initialized SimpleDocxConverter (local extraction)")

    async def convert_batch(self, paths: List[Path]) -> Dict[Path, str]:
        """Extract text from multiple DOCX files.

        Args:
            paths: List of DOCX file paths

        Returns:
            Dict mapping path to extracted markdown text
        """
        self.logger.debug(f"Converting {len(paths)} DOCX files")

        semaphore = asyncio.Semaphore(self.MAX_CONCURRENT)
        tasks = [self._convert_single(path, semaphore) for path in paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        output = {}
        for path, result in zip(paths, results):
            if isinstance(result, Exception):
                self.logger.error(f"Failed to convert {path.name}: {result}")
                output[path] = None
            else:
                output[path] = result

        success_count = sum(1 for v in output.values() if v is not None)
        self.logger.info(
            f"Converted {success_count}/{len(paths)} DOCX files successfully"
        )

        return output

    async def _convert_single(
        self, path: Path, semaphore: asyncio.Semaphore
    ) -> str:
        """Extract text from a single DOCX file.

        Args:
            path: Path to DOCX file
            semaphore: Concurrency control

        Returns:
            Extracted text as markdown
        """
        async with semaphore:
            try:
                def extract():
                    doc = Document(str(path))
                    paragraphs = []

                    # Extract paragraphs
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            # Preserve heading styles
                            if para.style.name.startswith('Heading'):
                                level = para.style.name.replace('Heading ', '')
                                try:
                                    level_num = int(level)
                                    paragraphs.append(f"{'#' * level_num} {text}")
                                except ValueError:
                                    paragraphs.append(text)
                            else:
                                paragraphs.append(text)

                    # Extract tables
                    for table in doc.tables:
                        table_rows = []
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            table_rows.append(" | ".join(cells))

                        if table_rows:
                            paragraphs.append("\n**Table:**")
                            paragraphs.append(table_rows[0])  # Header
                            if len(table_rows) > 1:
                                paragraphs.append(
                                    " | ".join(["---"] * len(table_rows[0].split(" | ")))
                                )
                                paragraphs.extend(table_rows[1:])  # Data rows

                    return "\n\n".join(paragraphs)

                text = await run_in_thread_pool(extract)

                if not text or not text.strip():
                    self.logger.warning(f"No text extracted from {path.name}")
                    return ""

                return text.strip()

            except Exception as e:
                self.logger.error(
                    f"Error extracting text from {path.name}: {e}",
                    exc_info=True
                )
                raise EntityProcessingError(
                    f"Failed to extract text from DOCX: {e}"
                )
